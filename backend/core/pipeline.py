import requests
from requests.auth import HTTPBasicAuth
from bs4 import BeautifulSoup
import os
from pinecone import Pinecone
import uuid
from sentence_transformers import SentenceTransformer
import io
from pypdf import PdfReader
from docx import Document
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.config import *

class KnowledgeBase:
    def __init__(self):
        self.documents = []
        self.pc = Pinecone(api_key=PINECONE_API_KEY)
        self.index = self.pc.Index(INDEX_NAME)
        # Load local embedding model
        print("Loading embedding model...")
        try:
            # Try using huggingface_hub directly to download
            from huggingface_hub import hf_hub_download, snapshot_download
            cache_dir = "./models_cache"
            os.makedirs(cache_dir, exist_ok=True)
            
            # Download the model snapshot first
            model_path = snapshot_download(
                repo_id="sentence-transformers/all-MiniLM-L6-v2",
                cache_dir=cache_dir,
                local_files_only=False
            )
            print(f"Model downloaded to: {model_path}")
            
            # Now load the model from the local path
            self.model = SentenceTransformer(model_path)
            print("Model loaded successfully!")
            
        except Exception as e:
            print(f"Error with HuggingFace approach: {e}")
            print("Using simple TF-IDF as fallback...")
            # Use sklearn's TfidfVectorizer as a simple fallback
            from sklearn.feature_extraction.text import TfidfVectorizer
            import numpy as np
            
            class SimpleEmbedder:
                def __init__(self):
                    self.vectorizer = TfidfVectorizer(max_features=384)  # Match embedding dimension
                
                def encode(self, texts):
                    if isinstance(texts, str):
                        texts = [texts]
                    embeddings = self.vectorizer.fit_transform(texts).toarray()
                    # Ensure we have exactly 384 dimensions
                    if embeddings.shape[1] < 384:
                        # Pad with zeros if needed
                        padding = 384 - embeddings.shape[1]
                        embeddings = np.pad(embeddings, ((0, 0), (0, padding)), 'constant')
                    return embeddings
                
                def __call__(self, texts):
                    return self.encode(texts)
            
            self.model = SimpleEmbedder()

    def clean_html(self, html_content):
        soup = BeautifulSoup(html_content, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text(separator=' ')
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return text

    def extract_text_from_pdf(self, pdf_bytes):
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_bytes):
        try:
            doc = Document(io.BytesIO(docx_bytes))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""
    
    def extract_text_from_html(self, html_bytes):
        try:
            html_content = html_bytes.decode('utf-8', errors='ignore')
            return self.clean_html(html_content)
        except Exception as e:
            print(f"Error reading HTML: {e}")
            return ""

    def get_attachments(self, page_id):
        url = f"https://{DOMAIN}/wiki/rest/api/content/{page_id}/child/attachment"
        response = requests.get(
            url,
            auth=HTTPBasicAuth(EMAIL, API_TOKEN),
            headers={"Accept": "application/json"}
        )
        if response.status_code == 200:
            return response.json().get('results', [])
        return []

    def fetch_confluence_page(self, page_id):
        # 1. Fetch Page Body
        url = f"https://{DOMAIN}/wiki/rest/api/content/{page_id}?expand=body.storage"
        print(f"Fetching Confluence Page ID: {page_id}...")
        response = requests.get(
            url,
            auth=HTTPBasicAuth(EMAIL, API_TOKEN),
            headers={"Accept": "application/json"}
        )
        if response.status_code != 200:
            print(f"Error fetching page {page_id}: {response.status_code}")
            return None

        data = response.json()
        title = data["title"]
        raw_html = data["body"]["storage"]["value"]
        clean_text = self.clean_html(raw_html)
        
        print(f"Successfully fetched page: {title}")
        print(f" - Page Text Length: {len(clean_text)} chars")

        # 2. Fetch Attachments
        print(f"Checking for attachments on page {page_id}...")
        attachments = self.get_attachments(page_id)
        attachment_text = ""
        
        for att in attachments:
            media_type = att['metadata']['mediaType']
            filename = att['title']
            download_link = att['_links']['download']
            download_url = f"https://{DOMAIN}/wiki{download_link}"
            
            print(f" - Found attachment: {filename} (Type: {media_type})")
            
            # Download the file
            try:
                file_response = requests.get(
                    download_url,
                    auth=HTTPBasicAuth(EMAIL, API_TOKEN)
                )
                
                if file_response.status_code != 200:
                    print(f"   Failed to download: {file_response.status_code}")
                    continue
                
                # Process based on file type
                content = ""
                if media_type == 'application/pdf':
                    content = self.extract_text_from_pdf(file_response.content)
                elif media_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                                   'application/msword']:
                    content = self.extract_text_from_docx(file_response.content)
                elif media_type in ['text/html', 'application/xhtml+xml']:
                    content = self.extract_text_from_html(file_response.content)
                else:
                    print(f"   Unsupported file type: {media_type}")
                    continue
                
                if content:
                    print(f"   Extracted {len(content)} chars from {filename}")
                    attachment_text += f"\n\n--- Attachment: {filename} ---\n{content}"
                else:
                    print(f"   No content extracted from {filename}")
                    
            except Exception as e:
                print(f"   Error processing {filename}: {e}")

        # Combine content
        full_content = clean_text + attachment_text
        print(f"Total Content Length to Index: {len(full_content)} chars")
        
        return {"source": f"Confluence - {title}", "content": full_content}

    def chunk_text(self, text, chunk_size=1000, overlap=100):
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start += chunk_size - overlap
        return chunks

    def upload_to_pinecone(self):
        print("\n--- UPLOADING TO PINECONE ---")
        records = []
        
        for doc in self.documents:
            print(f"Processing: {doc['source']}")
            chunks = self.chunk_text(doc['content'])
            print(f" - Content Length: {len(doc['content'])} chars")
            print(f" - Generated {len(chunks)} chunks")
            
            if len(chunks) == 0:
                print(" ⚠️ WARNING: No chunks generated! Check content.")
                continue

            # Generate embeddings for all chunks at once
            embeddings = self.model.encode(chunks)
            
            for i, chunk in enumerate(chunks):
                record_id = str(uuid.uuid4())
                
                # Prepare the record
                record = {
                    "id": record_id,
                    "values": embeddings[i].tolist(), # Convert numpy array to list
                    "metadata": {
                        "chunk_text": chunk,
                        "source": doc['source'],
                        "chunk_index": i
                    }
                }
                records.append(record)

        # Upload in batches
        batch_size = 50
        for i in range(0, len(records), batch_size):
            batch = records[i:i+batch_size]
            print(f"Upserting batch {i//batch_size + 1}...")
            try:
                self.index.upsert(vectors=batch)
            except Exception as e:
                print(f"Error upserting batch: {e}")

        print("--- UPLOAD COMPLETE ---")

    def run_pipeline(self):
        print("--- STARTING PIPELINE ---")
        # 1. Fetch Data (Confluence ONLY)
        for pid in PAGE_IDS:
            doc = self.fetch_confluence_page(pid)
            if doc:
                self.documents.append(doc)

        # 2. Upload to Pinecone
        self.upload_to_pinecone()

if __name__ == "__main__":
    kb = KnowledgeBase()
    kb.run_pipeline()
